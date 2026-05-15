"""
Resume Feature Extractor
Extracts structured features from raw resume text like a recruiter would
"""

import re
import logging
import warnings
import datetime
from pathlib import Path
from typing import List, Dict

import PyPDF2
from docx import Document
from PIL import Image
import pytesseract

# Set Tesseract path (Windows default; skipped if not found)
import shutil
_tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if Path(_tesseract_path).exists():
    pytesseract.pytesseract.tesseract_cmd = _tesseract_path
elif shutil.which('tesseract'):
    pytesseract.pytesseract.tesseract_cmd = shutil.which('tesseract')

from config import (
    TECHNICAL_SKILLS, SOFT_SKILLS, EDUCATION_LEVELS,
    CERTIFICATIONS,
    SENIORITY_LEVELS, QUALITY_CRITERIA, RED_FLAG_RULES,
)


class ResumeFeatureExtractor:
    """Extract structured features from resume text"""

    def __init__(self):
        self.technical_skills = [s.lower() for s in TECHNICAL_SKILLS]
        self.soft_skills = [s.lower() for s in SOFT_SKILLS]
        self.certifications = [c.lower() for c in CERTIFICATIONS]

        # Pre-compile Regex for faster skill matching (Crucial for large datasets)
        # We sort by length descending to match 'c++' before 'c', etc.
        sorted_tech = sorted(self.technical_skills, key=len, reverse=True)
        self.tech_regex = re.compile(r'\b(' + '|'.join(map(re.escape, sorted_tech)) + r')\b')
        
        sorted_soft = sorted(self.soft_skills, key=len, reverse=True)
        self.soft_regex = re.compile(r'\b(' + '|'.join(map(re.escape, sorted_soft)) + r')\b')

    def clean_text(self, text: str) -> str:
        """Clean and normalize text for better feature extraction and semantic matching."""
        if not text:
            return ""
        # 1. Remove non-printable characters
        text = "".join(char for char in text if char.isprintable() or char in "\n\t")
        # 2. Normalize whitespace (tabs to spaces, multiple spaces to one)
        text = re.sub(r'[ \t]+', ' ', text)
        # 3. Normalize newlines (limit to max 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 4. Strip leading/trailing whitespace
        return text.strip()

    def extract_all(self, text: str) -> Dict:
        """Extract ALL features from resume text — full ATS-grade extraction"""
        text = self.clean_text(text)
        text_lower = text.lower()

        features = {
            "raw_text": text,
            "text_length": len(text),
            "word_count": len(text.split()),

            # --- Candidate identity ---
            "candidate_name": self.extract_candidate_name(text),

            # --- Skills ---
            "technical_skills": self.extract_technical_skills(text_lower),
            "soft_skills": self.extract_soft_skills(text_lower),
            "num_technical_skills": 0,
            "num_soft_skills": 0,
            "total_skills": 0,

            # --- Experience ---
            "years_of_experience": self.extract_years_of_experience(text_lower),

            # --- Education ---
            "education_level": self.extract_education_level(text_lower),
            "education_score": self.extract_education_score(text_lower),

            # --- Certifications ---
            "certifications": self.extract_certifications(text_lower),
            "num_certifications": 0,

            # --- Contact info ---
            "has_email": self.has_email(text),
            "has_phone": self.has_phone(text),
            "email": self.extract_email(text),
            "phone": self.extract_phone(text),

            # --- Job titles ---
            "job_titles": self.extract_job_titles(text_lower),

            # --- Sections ---
            "num_sections": self.count_sections(text_lower),

            # --- Seniority (ATS) ---
            "seniority_level": "unknown",
            "seniority_score": 0,

            # --- Employment gaps (ATS) ---
            "employment_gaps": [],
            "num_employment_gaps": 0,
            "total_gap_months": 0,

            # --- Resume quality (ATS) ---
            "resume_quality_score": 0,
            "quality_breakdown": {},

            # --- Red flags (ATS) ---
            "red_flags": [],
            "num_red_flags": 0,
        }

        # Compute counts
        features["num_technical_skills"] = len(features["technical_skills"])
        features["num_soft_skills"] = len(features["soft_skills"])
        features["total_skills"] = features["num_technical_skills"] + features["num_soft_skills"]
        features["num_certifications"] = len(features["certifications"])

        # --- Seniority detection ---
        seniority = self.detect_seniority_level(
            text_lower, features["years_of_experience"], features["job_titles"]
        )
        features["seniority_level"] = seniority["level"]
        features["seniority_score"] = seniority["score"]

        # --- Employment gap detection ---
        gaps = self.detect_employment_gaps(text_lower)
        features["employment_gaps"] = gaps
        features["num_employment_gaps"] = len(gaps)
        features["total_gap_months"] = sum(g["months"] for g in gaps)

        # --- Resume quality scoring ---
        quality = self.compute_resume_quality_score(features)
        features["resume_quality_score"] = quality["score"]
        features["quality_breakdown"] = quality["breakdown"]

        # --- Red flag detection ---
        red_flags = self.detect_red_flags(features)
        features["red_flags"] = red_flags
        features["num_red_flags"] = len(red_flags)

        return features

    def extract_technical_skills(self, text: str) -> List[str]:
        """Extract technical skills found in text (Optimized)"""
        return list(set(self.tech_regex.findall(text)))

    def extract_soft_skills(self, text: str) -> List[str]:
        """Extract soft skills found in text (Optimized)"""
        return list(set(self.soft_regex.findall(text)))

    def extract_years_of_experience(self, text: str) -> int:
        """Extract years of experience from resume text"""
        # Direct patterns: "X years of experience"
        patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp(?:erience)?)',
            r'(?:experience|exp(?:erience)?)\s*(?:of\s*)?(\d+)\+?\s*(?:years?|yrs?)',
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:working|professional|industry)',
            r'over\s*(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
        ]

        all_years = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for m in matches:
                yr = int(m)
                if 1 <= yr <= 40:  # Realistic work experience range
                    all_years.append(yr)

        if all_years:
            return max(all_years)

        # Fallback: count work date ranges (only near "experience" or "work" context)
        # Look for patterns like "2015 - 2020" or "2015-present"
        work_ranges = re.findall(
            r'(20[0-2]\d|19[89]\d)\s*[-–to]+\s*(20[0-2]\d|present|current|till date|ongoing)',
            text
        )
        if work_ranges:
            current_year = datetime.datetime.now().year
            max_exp = 0
            for start, end in work_ranges:
                start_yr = int(start)
                end_yr = current_year if end in ('present', 'current', 'till date', 'ongoing') else int(end)
                diff = end_yr - start_yr
                if 0 < diff <= 40:
                    max_exp = max(max_exp, diff)
            if max_exp > 0:
                return max_exp

        return 0

    def extract_education_level(self, text: str) -> str:
        """Extract highest education level"""
        best_level = "unknown"
        best_score = 0

        for keyword, score in EDUCATION_LEVELS.items():
            # Use word boundary for short keywords to avoid false matches
            if len(keyword) <= 3:
                if re.search(r'\b' + re.escape(keyword) + r'\b', text):
                    if score > best_score:
                        best_score = score
                        best_level = keyword
            else:
                if keyword in text and score > best_score:
                    best_score = score
                    best_level = keyword

        # Map to clean label
        level_map = {
            6: "phd", 5: "masters", 4: "bachelors",
            3: "diploma", 2: "hsc", 1: "ssc", 0: "unknown"
        }
        return level_map.get(best_score, "unknown")

    def extract_education_score(self, text: str) -> int:
        """Return numeric education score (0-6)"""
        best_score = 0
        for keyword, score in EDUCATION_LEVELS.items():
            if len(keyword) <= 3:
                if re.search(r'\b' + re.escape(keyword) + r'\b', text):
                    if score > best_score:
                        best_score = score
            else:
                if keyword in text and score > best_score:
                    best_score = score
        return best_score

    def extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        found = []
        for cert in self.certifications:
            if cert in text:
                found.append(cert)
        return list(set(found))

    def has_email(self, text: str) -> bool:
        """Check if resume has email"""
        return bool(re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text))

    def extract_email(self, text: str) -> str:
        """Extract first email address"""
        match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        return match.group(0) if match else ""

    def has_phone(self, text: str) -> bool:
        """Check if resume has phone number"""
        return bool(re.search(r'(?<!\d)(?:\+?\d[\d\s\-\(\)]{8,}\d)(?!\d)', text))

    def extract_phone(self, text: str) -> str:
        """Extract first phone number"""
        match = re.search(r'(?<!\d)(?:\+?\d[\d\s\-\(\)]{8,}\d)(?!\d)', text)
        return match.group(0).strip() if match else ""

    def extract_job_titles(self, text: str) -> List[str]:
        """Extract common job titles"""
        titles = [
            "manager", "director", "engineer", "developer", "analyst",
            "accountant", "auditor", "officer", "executive", "assistant",
            "coordinator", "supervisor", "specialist", "consultant",
            "researcher", "scientist", "professor", "lecturer",
            "administrator", "secretary", "receptionist", "clerk",
            "technician", "operator", "helper", "peon", "attendant",
            "intern", "trainee", "lead", "head", "chief",
        ]
        found = [t for t in titles if t in text]
        return list(set(found))

    def count_sections(self, text: str) -> int:
        """Count resume sections (education, experience, skills, etc.)"""
        sections = [
            "education", "experience", "skill", "objective", "summary",
            "project", "certification", "achievement", "reference",
            "personal", "contact", "language", "hobby", "interest",
            "training", "award", "publication",
        ]
        return sum(1 for s in sections if s in text)

    # =================================================================
    # ATS-grade extraction (seniority, gaps, quality, red-flags, name)
    # =================================================================

    def extract_candidate_name(self, text: str) -> str:
        """Best-effort extraction of candidate name from resume header.
        Looks at the first few non-empty lines before any section heading."""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        section_words = {
            "education", "experience", "skill", "objective", "summary",
            "project", "certification", "profile", "contact", "address",
            "career", "personal", "about", "phone", "email",
        }
        for line in lines[:5]:
            low = line.lower()
            # Skip lines that look like sections, emails, phones, urls
            if any(sw in low for sw in section_words):
                continue
            if "@" in line or re.search(r'\d{7,}', line):
                continue
            if re.match(r'^https?://', line):
                continue
            # A name line is typically 2-4 capitalized words, no digits
            clean = re.sub(r'[^a-zA-Z\s\.]', '', line).strip()
            words = clean.split()
            if 2 <= len(words) <= 5 and all(w[0].isupper() or w == '.' for w in words if w):
                return clean
        return ""

    def detect_seniority_level(self, text: str, years_exp: int,
                                job_titles: List[str]) -> Dict:
        """Detect candidate seniority: entry/junior/mid/senior/director/executive."""
        best_level = "entry"
        best_score = 1

        # Method 1: Check title keywords in full text
        for level, info in SENIORITY_LEVELS.items():
            for title_kw in info["titles"]:
                if title_kw in text:
                    if info["score"] > best_score:
                        best_score = info["score"]
                        best_level = level
                    break  # one match per level is enough

        # Method 2: Infer from years of experience (if title didn't give strong signal)
        if best_score <= 2 and years_exp > 0:
            if years_exp >= 15:
                best_level, best_score = "executive", 6
            elif years_exp >= 10:
                best_level, best_score = "director", 5
            elif years_exp >= 6:
                best_level, best_score = "senior", 4
            elif years_exp >= 3:
                best_level, best_score = "mid", 3
            elif years_exp >= 1:
                best_level, best_score = "junior", 2

        return {"level": best_level, "score": best_score}

    def detect_employment_gaps(self, text: str) -> List[Dict]:
        """Detect gaps between employment periods.
        Returns list of {start_year, end_year, months} dicts."""
        current_year = datetime.datetime.now().year

        # Find all date ranges (e.g. "2015 - 2018", "Jan 2020 - Present")
        range_pattern = re.compile(
            r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?'
            r'[a-z]*\.?\s*(\d{4})\s*[-\x96\x97to]+\s*'
            r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?'
            r'[a-z]*\.?\s*(\d{4}|present|current|till\s*date|ongoing|now)',
            re.IGNORECASE
        )

        periods = []
        for m in range_pattern.finditer(text):
            start = int(m.group(1))
            end_raw = m.group(2).strip().lower()
            end = current_year if end_raw in ('present', 'current', 'till date',
                                               'ongoing', 'now') else int(end_raw)
            if 1980 <= start <= current_year and start <= end <= current_year + 1:
                periods.append((start, end))

        if len(periods) < 2:
            return []

        # Sort by start year and find gaps
        periods.sort(key=lambda x: x[0])
        gaps = []
        for i in range(len(periods) - 1):
            prev_end = periods[i][1]
            next_start = periods[i + 1][0]
            gap_years = next_start - prev_end
            if gap_years >= 1:  # Only flag gaps of 1+ year
                gaps.append({
                    "after_year": prev_end,
                    "before_year": next_start,
                    "months": gap_years * 12,
                })

        return gaps

    def compute_resume_quality_score(self, features: Dict) -> Dict:
        """Score resume completeness & professionalism (0-100)."""
        breakdown = {}
        score = 0

        # Name
        has_name = bool(features.get("candidate_name"))
        pts = QUALITY_CRITERIA["has_name"] if has_name else 0
        breakdown["has_name"] = pts
        score += pts

        # Email
        pts = QUALITY_CRITERIA["has_email"] if features.get("has_email") else 0
        breakdown["has_email"] = pts
        score += pts

        # Phone
        pts = QUALITY_CRITERIA["has_phone"] if features.get("has_phone") else 0
        breakdown["has_phone"] = pts
        score += pts

        # Education
        has_edu = features.get("education_score", 0) > 0
        pts = QUALITY_CRITERIA["has_education"] if has_edu else 0
        breakdown["has_education"] = pts
        score += pts

        # Experience
        has_exp = features.get("years_of_experience", 0) > 0
        pts = QUALITY_CRITERIA["has_experience"] if has_exp else 0
        breakdown["has_experience"] = pts
        score += pts

        # Skills
        has_skills = features.get("total_skills", 0) >= 2
        pts = QUALITY_CRITERIA["has_skills"] if has_skills else 0
        breakdown["has_skills"] = pts
        score += pts

        # Objective / summary
        raw = features.get("raw_text", "").lower()
        has_obj = any(kw in raw for kw in ["objective", "summary", "profile", "about me"])
        pts = QUALITY_CRITERIA["has_objective"] if has_obj else 0
        breakdown["has_objective"] = pts
        score += pts

        # Certifications
        pts = QUALITY_CRITERIA["has_certifications"] if features.get("num_certifications", 0) > 0 else 0
        breakdown["has_certifications"] = pts
        score += pts

        # Word count (not too short, not too long)
        wc = features.get("word_count", 0)
        word_ok = 100 <= wc <= 5000
        pts = QUALITY_CRITERIA["word_count_ok"] if word_ok else (QUALITY_CRITERIA["word_count_ok"] // 2 if wc > 50 else 0)
        breakdown["word_count_ok"] = pts
        score += pts

        # Multiple sections
        pts = QUALITY_CRITERIA["multiple_sections"] if features.get("num_sections", 0) >= 3 else 0
        breakdown["multiple_sections"] = pts
        score += pts

        # Recent experience (has dates within last 3 years)
        cur = datetime.datetime.now().year
        recent_years = re.findall(r'20[12]\d', raw)
        has_recent = any(cur - int(y) <= 3 for y in recent_years) if recent_years else False
        pts = QUALITY_CRITERIA["recent_experience"] if has_recent else 0
        breakdown["recent_experience"] = pts
        score += pts

        # Red flags placeholder (computed after this)
        breakdown["no_red_flags"] = QUALITY_CRITERIA["no_red_flags"]
        score += QUALITY_CRITERIA["no_red_flags"]

        return {"score": min(score, 100), "breakdown": breakdown}

    def detect_red_flags(self, features: Dict) -> List[Dict]:
        """Detect red flags that a recruiter would notice."""
        flags = []

        # 1. No contact info
        if not features.get("has_email") and not features.get("has_phone"):
            flags.append({
                "flag": "no_contact",
                "description": RED_FLAG_RULES["no_contact"]["description"],
                "severity": RED_FLAG_RULES["no_contact"]["severity"],
            })

        # 2. Very short resume
        if features.get("word_count", 0) < 100:
            flags.append({
                "flag": "very_short",
                "description": RED_FLAG_RULES["very_short"]["description"],
                "severity": RED_FLAG_RULES["very_short"]["severity"],
            })

        # 3. No experience
        if features.get("years_of_experience", 0) == 0 and features.get("seniority_score", 0) <= 1:
            flags.append({
                "flag": "no_experience",
                "description": RED_FLAG_RULES["no_experience"]["description"],
                "severity": RED_FLAG_RULES["no_experience"]["severity"],
            })

        # 4. Long employment gap
        for gap in features.get("employment_gaps", []):
            if gap.get("months", 0) >= 24:
                flags.append({
                    "flag": "long_gap",
                    "description": "{} ({} - {})".format(
                        RED_FLAG_RULES["long_gap"]["description"],
                        gap["after_year"], gap["before_year"]
                    ),
                    "severity": RED_FLAG_RULES["long_gap"]["severity"],
                })

        # 5. Job hopping (many short stints)
        raw = features.get("raw_text", "").lower()
        cur_year = datetime.datetime.now().year
        range_pat = re.compile(
            r'(\d{4})\s*[-\x96\x97to]+\s*(\d{4}|present|current)'
        )
        stints = []
        for m in range_pat.finditer(raw):
            s = int(m.group(1))
            e_raw = m.group(2).lower()
            e = cur_year if e_raw in ('present', 'current') else int(e_raw)
            if 1980 <= s <= cur_year:
                stints.append(e - s)
        short_stints = [d for d in stints if 0 < d <= 1]
        if len(short_stints) >= 3:
            flags.append({
                "flag": "job_hopping",
                "description": RED_FLAG_RULES["job_hopping"]["description"],
                "severity": RED_FLAG_RULES["job_hopping"]["severity"],
            })

        # Deduct quality points if red flags found
        if flags and features.get("resume_quality_score", 0) > 0:
            penalty = min(len(flags) * 3, QUALITY_CRITERIA["no_red_flags"])
            features["resume_quality_score"] = max(
                0, features["resume_quality_score"] - penalty
            )
            if "quality_breakdown" in features:
                features["quality_breakdown"]["no_red_flags"] = max(
                    0, QUALITY_CRITERIA["no_red_flags"] - penalty
                )

        return flags


def extract_text_from_file(file_path: Path) -> str:
    """Extract text from PDF, DOCX, DOC, TXT, or image files"""
    suffix = file_path.suffix.lower()

    try:
        if suffix == '.pdf':
            text = ""
            # Suppress all PyPDF2 warnings
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                logging.getLogger("PyPDF2").setLevel(logging.CRITICAL)
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f, strict=False)
                        for page in reader.pages:
                            try:
                                page_text = page.extract_text() or ""
                                # Remove surrogate characters that cause encoding errors
                                page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                                text += page_text
                            except (KeyboardInterrupt, SystemExit):
                                raise
                            except Exception:
                                continue
                except (KeyboardInterrupt, SystemExit, RuntimeError, TypeError):
                    # Skip problematic PDFs (e.g., corrupted or incompatible with PyPDF2)
                    return ""
            
            text = re.sub(r'(?<! ) \n', '\n', text) # Clean up PyPDF2 artifacts
            
            if len(text.strip()) < 100:
                # Fallback: Treat PDF as image and use OCR if it is empty/scanned
                try:
                    from pdf2image import convert_from_path # type: ignore
                    images = convert_from_path(file_path)
                    text = ""
                    for img in images:
                        text += pytesseract.image_to_string(img)
                except (ImportError, Exception):
                    pass
            
            # Final cleanup
            text = "".join(char for char in text if char.isprintable() or char in "\n\t")
            text = re.sub(r'[ \t]+', ' ', text)
            return text.strip()

        elif suffix == '.docx':
            doc = Document(str(file_path))
            return "\n".join(para.text for para in doc.paragraphs).strip()

        elif suffix == '.doc':
            try:
                doc = Document(str(file_path))
                return "\n".join(para.text for para in doc.paragraphs).strip()
            except Exception:
                return ""

        elif suffix in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff'):
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()

        elif suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read().strip()

    except Exception:
        pass

    return ""